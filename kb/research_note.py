from __future__ import annotations

import io
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


_TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")
_INVALID_FILENAME_RE = re.compile(r'[\\/:*?"<>|\x00-\x1f]')
_INLINE_TOKEN_RE = re.compile(
    r"(!?\[[^\]]+\]\([^\s)]+(?:\s+[^)]*)?\)|\*\*[^*\n]+\*\*|`[^`\n]+`|(?<!\*)\*[^*\n]+\*(?!\*))"
)


def research_note_filename(title: str) -> str:
    """Return a readable Windows-safe DOCX filename for a research note."""

    cleaned = _INVALID_FILENAME_RE.sub("-", _clean_text(title, limit=120))
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .-")
    return f"{cleaned or 'Pi_zaya research note'}.docx"


def _clean_text(value: object, *, limit: int) -> str:
    text = str(value or "").replace("\x00", "").strip()
    return text[:limit]


def _split_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [cell.strip().replace(r"\|", "|") for cell in raw.split("|")]


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(_TABLE_SEPARATOR_CELL_RE.fullmatch(cell.replace(" ", "")) for cell in cells)


def _set_run_font(run, *, name: str = "Aptos", east_asia: str = "等线", size: float | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F6FAF")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Aptos")
    fonts.set(qn("w:hAnsi"), "Aptos")
    fonts.set(qn("w:eastAsia"), "等线")
    properties.append(fonts)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in _INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            _set_run_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        image_match = re.fullmatch(r"!\[([^\]]+)\]\(([^\s)]+)(?:\s+[^)]*)?\)", token)
        link_match = re.fullmatch(r"\[([^\]]+)\]\(([^\s)]+)(?:\s+[^)]*)?\)", token)
        if image_match:
            _set_run_font(paragraph.add_run(f"[Image: {image_match.group(1)}]"))
        elif link_match:
            _add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            _set_run_font(run)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name="Consolas", east_asia="等线")
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F4F7")
            run._element.get_or_add_rPr().append(shading)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            _set_run_font(run)
        else:
            _set_run_font(paragraph.add_run(token))
        cursor = match.end()
    if cursor < len(text):
        _set_run_font(paragraph.add_run(text[cursor:]))


def _style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.color.rgb = None


def _add_table(document: Document, rows: Iterable[list[str]]) -> None:
    table_rows = list(rows)
    if not table_rows:
        return
    column_count = max(len(row) for row in table_rows)
    table = document.add_table(rows=len(table_rows), cols=column_count)
    table.style = "Light Shading Accent 1"
    for row_index, source_row in enumerate(table_rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            value = source_row[column_index] if column_index < len(source_row) else ""
            paragraph = cell.paragraphs[0]
            _add_inline_markdown(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def research_note_docx(title: str, content_markdown: str) -> bytes:
    """Render user-edited research-note Markdown to a real, portable DOCX file."""

    safe_title = _clean_text(title, limit=240) or "Pi_zaya research note"
    content = _clean_text(content_markdown, limit=160_000)
    document = Document()
    _style_document(document)
    document.core_properties.title = safe_title
    document.core_properties.subject = "Pi_zaya research note"
    document.add_heading(safe_title, level=0)

    lines = content.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Normal"]
                run = paragraph.add_run("\n".join(code_lines))
                _set_run_font(run, name="Consolas", east_asia="等线", size=9)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F5F6F8")
                paragraph._p.get_or_add_pPr().append(shading)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw_line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if (
            "|" in stripped
            and index + 1 < len(lines)
            and _is_table_separator(lines[index + 1])
        ):
            table_rows = [_split_table_row(stripped)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_rows.append(_split_table_row(lines[index]))
                index += 1
            _add_table(document, table_rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            paragraph = document.add_heading(level=min(4, len(heading.group(1))))
            _add_inline_markdown(paragraph, heading.group(2).strip())
            index += 1
            continue
        if re.match(r"^[-*+]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_markdown(paragraph, re.sub(r"^[-*+]\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_markdown(paragraph, re.sub(r"^\d+[.)]\s+", "", stripped))
            index += 1
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.55)
            border = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "10")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "A8C7E6")
            border.append(left)
            paragraph._p.get_or_add_pPr().append(border)
            _add_inline_markdown(paragraph, stripped.lstrip("> "))
            index += 1
            continue
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            index += 1
            continue
        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, stripped)
        index += 1

    if in_code and code_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run("\n".join(code_lines))
        _set_run_font(run, name="Consolas", east_asia="等线", size=9)

    for section in document.sections:
        if section.start_type == WD_SECTION.NEW_PAGE:
            section.start_type = WD_SECTION.CONTINUOUS
        footer = section.footer.paragraphs[0]
        footer.alignment = 2
        run = footer.add_run("Pi_zaya")
        _set_run_font(run, size=8)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
