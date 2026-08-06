from __future__ import annotations

import io
import re
import time
from pathlib import Path
from typing import Any


_MAX_BRIEF_SOURCES = 8
_MAX_EVIDENCE_TEXT = 1_800
_CITATION_RE = re.compile(r"(?<!\[)\[(\d+(?:\s*(?:,|-)\s*\d+)*)\](?!\])")
_MATRIX_UNVERIFIABLE_BOUNDARY_RE = re.compile(
    r"\b(?:not|never)\s+(?:directly\s+)?(?:comparable|compared|reported|stated|provided|available)\b"
    r"|\bno\s+(?:retrieved|available|selected)\s+(?:snippet|evidence|source)"
    r"|(?:不可直接比较|未报告|未说明|未提供|没有可用证据)",
    flags=re.IGNORECASE,
)


def _text(value: object, *, limit: int = 2_000) -> str:
    if value is None:
        return ""
    text = str(value).replace("\x00", " ").strip()
    return text[: max(0, int(limit))]


def _first_text(item: dict[str, Any], *keys: str, limit: int = 2_000) -> str:
    for key in keys:
        value = _text(item.get(key), limit=limit)
        if value:
            return value
    return ""


def _safe_int(value: object, *, default: int = 0) -> int:
    try:
        return int(value or 0)
    except (TypeError, ValueError, OverflowError):
        return int(default)


def _safe_float(value: object, *, default: float = 0.0) -> float:
    try:
        number = float(value or 0.0)
    except (TypeError, ValueError, OverflowError):
        return float(default)
    return number if number == number and abs(number) != float("inf") else float(default)


def _citation_numbers(answer: object, *, limit: int = 100) -> list[int]:
    numbers: list[int] = []
    seen: set[int] = set()
    for match in _CITATION_RE.finditer(str(answer or "")):
        for part in re.split(r"\s*,\s*", match.group(1)):
            raw = part.strip()
            if "-" in raw:
                start_raw, end_raw = [item.strip() for item in raw.split("-", 1)]
                start = _safe_int(start_raw)
                end = _safe_int(end_raw)
                if start <= 0 or end < start or end - start >= limit:
                    continue
                candidates = range(start, end + 1)
            else:
                candidates = (_safe_int(raw),)
            for number in candidates:
                if number <= 0 or number in seen:
                    continue
                seen.add(number)
                numbers.append(number)
                if len(numbers) >= limit:
                    return numbers
    return numbers


def _source_path(item: dict[str, Any]) -> str:
    library_path = _first_text(
        item,
        "libraryMatchPath",
        "library_match_path",
        limit=1_200,
    )
    if library_path:
        return library_path
    kind = _first_text(item, "shelfItemKind", "shelf_item_kind", limit=80).lower()
    route = _first_text(item, "citationRoute", "citation_route", limit=80).lower()
    if kind in {"reference", "inpaper", "reader_reference", "reader_references"}:
        return ""
    if route == "system_b" or item.get("isInpaper") is True or item.get("is_inpaper") is True:
        return ""
    return _first_text(
        item,
        "sourcePath",
        "source_path",
        limit=1_200,
    )


def _source_name(item: dict[str, Any]) -> str:
    return _first_text(
        item,
        "libraryMatchTitle",
        "library_match_title",
        "title",
        "cardTitle",
        "main",
        "sourceName",
        "source_name",
        limit=500,
    )


def _source_variants(value: object) -> set[str]:
    raw = _text(value, limit=1_200).replace("\\", "/")
    if not raw:
        return set()
    variants = {raw.lower()}
    try:
        path = Path(raw)
        variants.add(path.name.lower())
        variants.add(path.stem.lower())
        variants.add(str(path.expanduser().resolve(strict=False)).replace("\\", "/").lower())
    except Exception:
        pass
    return {item for item in variants if item}


def select_research_brief_sources(
    shelf_items: list[dict[str, Any]],
    *,
    item_keys: list[str] | None = None,
    limit: int = _MAX_BRIEF_SOURCES,
) -> list[dict[str, Any]]:
    requested = {_text(key, limit=500) for key in list(item_keys or []) if _text(key, limit=500)}
    selected: list[dict[str, Any]] = []
    seen: set[str] = set()
    for raw in list(shelf_items or []):
        if not isinstance(raw, dict):
            continue
        key = _first_text(raw, "key", "id", limit=500)
        if requested and key not in requested:
            continue
        source_path = _source_path(raw)
        source_name = _source_name(raw)
        if not source_path or not source_name:
            continue
        identity = source_path.replace("\\", "/").lower()
        if identity in seen:
            continue
        seen.add(identity)
        selected.append(dict(raw))
        if len(selected) >= max(1, min(_MAX_BRIEF_SOURCES, int(limit or _MAX_BRIEF_SOURCES))):
            break
    return selected


def research_brief_context(
    selected_items: list[dict[str, Any]],
    *,
    conversation_id: str = "",
) -> dict[str, Any]:
    items: list[dict[str, Any]] = []
    for index, item in enumerate(list(selected_items or []), start=1):
        source_path = _source_path(item)
        source_name = _source_name(item)
        title = _first_text(item, "title", "cardTitle", "main", limit=500) or source_name
        excerpt = _first_text(
            item,
            "shelfExcerpt",
            "shelf_excerpt",
            "evidenceQuote",
            "evidence_quote",
            "cardEvidence",
            "card_evidence",
            "citationContext",
            "citation_context",
            limit=1_600,
        )
        summary = _first_text(
            item,
            "summaryLine",
            "summary_line",
            "cardTakeaway",
            "card_takeaway",
            "whyLine",
            "why_line",
            limit=1_000,
        )
        items.append(
            {
                "key": _first_text(item, "key", "id", limit=500) or f"brief-source-{index}",
                "title": title,
                "sourceName": source_name,
                "sourcePath": source_path,
                "locationLabel": _first_text(item, "locationLabel", "location_label", limit=500),
                "headingPath": _first_text(item, "headingPath", "heading_path", limit=800),
                "blockId": _first_text(item, "blockId", "block_id", limit=500),
                "anchorId": _first_text(item, "anchorId", "anchor_id", "anchor", limit=500),
                "doi": _first_text(item, "doi", "doiUrl", "doi_url", limit=400),
                "authors": _first_text(item, "authors", limit=500),
                "year": _first_text(item, "year", limit=40),
                "summary": summary,
                "excerpt": excerpt,
                "note": _first_text(item, "note", limit=1_000),
            }
        )
    return {
        "version": 1,
        "id": f"research-brief:{_text(conversation_id, limit=120) or 'project'}",
        "source": "citation_shelf",
        "conversationId": _text(conversation_id, limit=120),
        "itemCount": len(items),
        "items": items,
    }


def research_brief_prompt(objective: str, *, locale: str = "zh") -> str:
    goal = _text(objective, limit=4_000)
    if str(locale or "").strip().lower().startswith("en"):
        return (
            "Create an evidence-grounded research brief using only the selected literature-basket sources. "
            f"Research objective: {goal or 'Synthesize the selected papers into an actionable research brief.'}\n\n"
            "Use these sections: Executive findings; Methods and experimental comparison; Quantitative evidence; "
            "Disagreements or boundary conditions; Research gaps and next questions. Write no more than eight substantive "
            "bullet sentences in total. Each bullet must be exactly one sentence and must end with its matching numeric "
            "citation such as [1] or [1, 2]. Headings must contain no factual claims. Use only citation numbers supplied "
            "with the evidence. Keep each paper's conditions attached to that paper; do not merge conditions from different "
            "papers. Omit any unsupported section or state a concise evidence gap without inventing background knowledge."
        )
    return (
        "仅使用已选文献篮中的本地论文证据，生成一份可继续编辑的研究简报。"
        f"研究目标：{goal or '综合所选论文，形成可执行的研究简报。'}\n\n"
        "请包含：核心结论；方法与实验条件对比；定量证据；分歧或适用边界；研究空白与下一步问题。"
        "全文最多写八条实质性要点；每条只能有一个完整句子，句末必须带匹配的数字引用，例如 [1] 或 [1, 2]。"
        "标题中不得写事实主张，只能使用证据中已有的引用编号。不同论文的条件必须分别绑定到原论文，不得拼接。"
        "没有证据支持的章节应省略，或只简洁说明证据缺口，不得用常识或外部知识补齐。"
    )


def research_brief_evidence(hits: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for index, hit in enumerate(list(hits or []), start=1):
        if not isinstance(hit, dict):
            continue
        meta = hit.get("meta") if isinstance(hit.get("meta"), dict) else {}
        source_path = _first_text(meta, "source_path", limit=1_200)
        source_name = _first_text(meta, "source_name", "title", limit=500)
        if not source_name and source_path:
            source_name = Path(source_path).stem
        out.append(
            {
                "citation_number": index,
                "source_path": source_path,
                "source_name": source_name,
                "title": _first_text(meta, "title", "source_name", limit=500),
                "heading_path": _first_text(meta, "heading_path", "top_heading", limit=800),
                "location_label": _first_text(meta, "location_label", "heading_path", limit=500),
                "page_start": meta.get("page_start") or meta.get("page") or None,
                "page_end": meta.get("page_end") or meta.get("page") or None,
                "block_id": _first_text(meta, "block_id", limit=500),
                "anchor_id": _first_text(meta, "anchor_id", "anchor", limit=500),
                "evidence_quote": _text(hit.get("text"), limit=_MAX_EVIDENCE_TEXT),
                "score": _safe_float(hit.get("score")),
            }
        )
    return out


def _item_matches_evidence(item: dict[str, Any], evidence: dict[str, Any]) -> bool:
    item_terms = _source_variants(_source_path(item)) | _source_variants(_source_name(item))
    evidence_terms = _source_variants(evidence.get("source_path")) | _source_variants(evidence.get("source_name"))
    return bool(item_terms & evidence_terms)


def research_brief_bibliography(
    selected_items: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: dict[str, int] = {}
    for evidence_item in list(evidence or []):
        match = next(
            (item for item in selected_items if _item_matches_evidence(item, evidence_item)),
            {},
        )
        doi = _first_text(match, "doi", "doiUrl", "doi_url", limit=400)
        source_path = _first_text(evidence_item, "source_path", limit=1_200) or _source_path(match)
        title = (
            _first_text(match, "title", "cardTitle", "main", limit=800)
            or _first_text(evidence_item, "title", "source_name", limit=800)
        )
        identity = (doi.lower() if doi else source_path.replace("\\", "/").lower()) or title.lower()
        citation_number = _safe_int(evidence_item.get("citation_number"))
        if identity in seen:
            row = out[seen[identity]]
            numbers = row.setdefault("citation_numbers", [])
            if citation_number > 0 and citation_number not in numbers:
                numbers.append(citation_number)
            continue
        seen[identity] = len(out)
        out.append(
            {
                "title": title,
                "authors": _first_text(match, "authors", limit=800),
                "year": _first_text(match, "year", limit=40),
                "venue": _first_text(match, "venue", limit=500),
                "doi": doi,
                "source_path": source_path,
                "source_name": _first_text(evidence_item, "source_name", limit=500) or _source_name(match),
                "citation_numbers": [citation_number] if citation_number > 0 else [],
            }
        )
    return out


def research_brief_quality(
    *,
    answer: str,
    agent_trace: dict[str, Any],
    selected_items: list[dict[str, Any]],
    evidence: list[dict[str, Any]],
) -> tuple[str, dict[str, Any]]:
    trace = agent_trace if isinstance(agent_trace, dict) else {}
    verification = trace.get("verification") if isinstance(trace.get("verification"), dict) else {}
    summary = trace.get("summary") if isinstance(trace.get("summary"), dict) else {}
    generation_quality_gate = str(summary.get("quality_gate_status") or "").strip().lower()
    claim_repair = summary.get("claim_repair") if isinstance(summary.get("claim_repair"), dict) else {}
    if generation_quality_gate == "fallback":
        generation_mode = "extractive_fallback"
    elif generation_quality_gate == "repaired" and bool(claim_repair.get("attempted")):
        generation_mode = "model_synthesis_repaired"
    else:
        generation_mode = "model_synthesis"
    total_claims = _safe_int(verification.get("total_claims") or summary.get("total_claims"))
    supported_claims = _safe_int(verification.get("supported_claims") or summary.get("supported_claims"))
    unsupported_claims = _safe_int(
        verification.get("unsupported_claims") or summary.get("unsupported_claims")
    )
    support_ratio = _safe_float(verification.get("support_ratio") or summary.get("support_ratio"))
    if total_claims > 0 and support_ratio <= 0:
        support_ratio = supported_claims / total_claims
    cited_numbers = sorted(_citation_numbers(answer))
    cited_evidence = [
        row for row in evidence if _safe_int(row.get("citation_number")) in cited_numbers
    ]
    expected_terms: set[str] = set()
    for item in selected_items:
        expected_terms.update(_source_variants(_source_path(item)))
        expected_terms.update(_source_variants(_source_name(item)))
    unexpected_sources: list[str] = []
    for row in evidence:
        variants = _source_variants(row.get("source_path")) | _source_variants(row.get("source_name"))
        if not variants or not (variants & expected_terms):
            unexpected_sources.append(_text(row.get("source_name") or row.get("source_path"), limit=500))
    selected_sources_without_evidence = [
        _source_name(item) or _source_path(item)
        for item in selected_items
        if not any(_item_matches_evidence(item, row) for row in cited_evidence)
    ]
    reasons: list[str] = []
    query_scope = str(summary.get("query_scope") or "").strip().lower()
    if query_scope != "basket":
        reasons.append("query_scope_not_basket")
    if generation_quality_gate not in {"passed", "repaired", "fallback"}:
        reasons.append("generation_quality_gate_missing")
    if str(trace.get("status") or "").strip().lower() != "done":
        reasons.append("agent_not_done")
    if list(trace.get("errors") or []):
        reasons.append("agent_errors")
    if str(verification.get("evidence_status") or summary.get("evidence_status") or "").strip().lower() != "grounded":
        reasons.append("evidence_not_grounded")
    if total_claims <= 0:
        reasons.append("no_audited_claims")
    if unsupported_claims > 0 or support_ratio < 0.999:
        reasons.append("unsupported_claims")
    if not evidence:
        reasons.append("no_evidence_hits")
    if unexpected_sources:
        reasons.append("unexpected_sources")
    if selected_sources_without_evidence:
        reasons.append("selected_sources_without_evidence")
    evidence_numbers = {_safe_int(row.get("citation_number")) for row in evidence}
    unresolved_citations = [number for number in cited_numbers if number not in evidence_numbers]
    if not cited_numbers:
        reasons.append("no_visible_citations")
    if unresolved_citations:
        reasons.append("unresolved_citations")
    quality_warnings = list(summary.get("quality_gate_warnings") or [])
    if generation_mode == "extractive_fallback":
        quality_warnings.append("extractive_fallback")
    quality = {
        "contract_version": 1,
        "query_scope": query_scope,
        "evidence_status": str(verification.get("evidence_status") or summary.get("evidence_status") or ""),
        "total_claims": total_claims,
        "supported_claims": supported_claims,
        "unsupported_claims": unsupported_claims,
        "support_ratio": round(max(0.0, min(1.0, support_ratio)), 4),
        "evidence_hit_count": len(evidence),
        "selected_source_count": len(selected_items),
        "citation_numbers": cited_numbers,
        "unresolved_citations": unresolved_citations,
        "unexpected_sources": [item for item in unexpected_sources if item],
        "selected_sources_without_evidence": [
            item for item in selected_sources_without_evidence if item
        ],
        "generation_quality_gate": generation_quality_gate,
        "generation_mode": generation_mode,
        "claim_repair": claim_repair,
        "warnings": list(dict.fromkeys(str(item) for item in quality_warnings if str(item or "").strip())),
        "reasons": reasons,
        "edited_after_verification": False,
    }
    return ("verified" if not reasons else "needs_review"), quality


def _matrix_source_key(item: dict[str, Any]) -> str:
    meta = item.get("meta") if isinstance(item.get("meta"), dict) else item
    return _first_text(meta, "source_path", "source_name", limit=1_200).replace("\\", "/").lower()


def _matrix_extractive_sentence(hit: dict[str, Any], citation_number: int) -> str:
    sentence = re.sub(r"\[[0-9][0-9,\-\s]*\]", " ", _text(hit.get("text"), limit=520))
    sentence = re.sub(r"[*_`|]+", " ", sentence)
    sentence = re.sub(r"\s+", " ", sentence).strip(" -:;,.。；：")
    return f"{sentence} [{citation_number}]." if sentence else ""


def _matrix_extractive_brief_answer(prompt: str, hits: list[dict[str, Any]]) -> str:
    prefer_zh = bool(re.search(r"[\u4e00-\u9fff]", str(prompt or "")))
    lines = ["## 证据" if prefer_zh else "## Evidence"]
    seen_sources: set[str] = set()
    for index, hit in enumerate(hits, start=1):
        if not isinstance(hit, dict):
            continue
        source_key = _matrix_source_key(hit)
        if not source_key or source_key in seen_sources:
            continue
        seen_sources.add(source_key)
        sentence = _matrix_extractive_sentence(hit, index)
        if sentence:
            lines.append(f"- {sentence}")
        if len(seen_sources) >= 8:
            break
    return "\n".join(lines).strip()


def _matrix_claim_plan(hits: list[dict[str, Any]], *, limit: int = 8) -> list[dict[str, Any]]:
    field_order = {
        "method": 0,
        "dataset_or_experiment": 1,
        "comparison_result": 2,
        "key_result": 3,
        "metric": 4,
        "limitation": 5,
    }
    grouped: dict[str, list[dict[str, Any]]] = {}
    source_order: list[str] = []
    for citation_number, hit in enumerate(hits, start=1):
        if not isinstance(hit, dict):
            continue
        source_key = _matrix_source_key(hit)
        if not source_key:
            continue
        if source_key not in grouped:
            grouped[source_key] = []
            source_order.append(source_key)
        meta = hit.get("meta") if isinstance(hit.get("meta"), dict) else {}
        evidence = re.sub(r"\[[0-9][0-9,\-\s]*\]", " ", _text(hit.get("text"), limit=420))
        evidence = re.sub(r"\s+", " ", evidence).strip()
        evidence_terms = {
            token.lower()
            for token in re.findall(r"[A-Za-z][A-Za-z0-9_-]{3,}|[\u4e00-\u9fff]{2,}", evidence)
            if token.lower() not in {"cdot", "times"}
        }
        if len(evidence_terms) < 2:
            continue
        grouped[source_key].append(
            {
                "citation": citation_number,
                "source_name": _first_text(meta, "source_name", "title", limit=180),
                "field": _first_text(meta, "matrix_field", limit=80),
                "evidence": evidence,
            }
        )
    for candidates in grouped.values():
        candidates.sort(key=lambda item: (field_order.get(str(item.get("field") or ""), 99), int(item["citation"])))
    plan: list[dict[str, Any]] = []
    depth = 0
    max_items = max(1, int(limit))
    while len(plan) < max_items:
        added = False
        for source_key in source_order:
            candidates = grouped[source_key]
            if depth < len(candidates):
                plan.append(candidates[depth])
                added = True
                if len(plan) >= max_items:
                    break
        if not added:
            break
        depth += 1
    return plan


def _matrix_generation_prompt(prompt: str, claim_plan: list[dict[str, Any]]) -> str:
    plan_lines = [
        f"- [{int(item.get('citation') or 0)}] {item.get('source_name') or 'source'} | "
        f"{item.get('field') or 'evidence'} | {item.get('evidence') or ''}"
        for item in claim_plan
    ]
    return (
        f"{prompt}\n\n"
        "Verified evidence-matrix brief contract:\n"
        "- Use only these optional section headings: Executive findings; Methods and experimental comparison; "
        "Quantitative evidence; Disagreements or boundary conditions. Do not add an overall title.\n"
        f"- Write exactly {len(claim_plan)} bullets: one bullet for each source-balanced claim-plan item below, in "
        "plan order. Do not add facts from snippets outside the plan or use a plan item twice.\n"
        "- Each bullet must be exactly one complete sentence and end with its supporting numeric citation marker. "
        "Spell out 'versus'; do not use the abbreviation 'vs.'.\n"
        "- Use exactly one numeric citation marker in each bullet. Compare sources through adjacent source-specific "
        "bullets rather than combining multiple sources in one sentence.\n"
        "- Name the method or paper associated with that citation. If its source name contains a publication year, "
        "use only that exact year and never transfer a year between sources.\n"
        "- Include at least one supported bullet for every source named in the claim plan. A cross-source sentence must "
        "cite evidence for every paper-specific clause.\n"
        "- Do not assert that evidence or a comparison is absent. Omit unsupported gaps, recommendations, and general "
        "background. Use only facts in the claim plan and retrieved snippets.\n\n"
        "Source-balanced claim plan:\n"
        + "\n".join(plan_lines)
    ).strip()


def _matrix_verified_sources(verification: dict[str, Any]) -> set[str]:
    sources: set[str] = set()
    for claim in list(verification.get("claims") or []):
        if not isinstance(claim, dict) or not bool(claim.get("supported")):
            continue
        for matched in list(claim.get("matched_sources") or []):
            if not isinstance(matched, dict):
                continue
            key = _first_text(matched, "source_path", "source_name", limit=1_200).replace("\\", "/").lower()
            if key:
                sources.add(key)
    return sources


def _matrix_claim_fully_bound(claim: dict[str, Any]) -> bool:
    claim_text = str(claim.get("claim_text") or claim.get("text") or "")
    if _MATRIX_UNVERIFIABLE_BOUNDARY_RE.search(claim_text):
        return False
    claim_years = set(re.findall(r"\b(?:19|20)\d{2}\b", claim_text))
    citation_numbers = {_safe_int(item) for item in list(claim.get("citation_numbers") or []) if _safe_int(item) > 0}
    if claim_years:
        matched_years = {
            year
            for item in list(claim.get("matched_sources") or [])
            if isinstance(item, dict)
            for year in re.findall(
                r"\b(?:19|20)\d{2}\b",
                f"{item.get('source_name') or ''} {item.get('source_path') or ''}",
            )
        }
        if not claim_years.issubset(matched_years):
            return False
    matched_numbers = {
        _safe_int(item.get("citation_index"))
        for item in list(claim.get("matched_sources") or [])
        if isinstance(item, dict) and _safe_int(item.get("citation_index")) > 0
    }
    if not bool(claim.get("supported") and citation_numbers and citation_numbers.issubset(matched_numbers)):
        return False
    contrast_clauses = [
        clause.strip()
        for clause in re.split(
            r"\s+(?:whereas|while|although|though|however|but)\s+|[;；]",
            claim_text,
            flags=re.IGNORECASE,
        )
        if clause.strip()
    ]
    if len(contrast_clauses) <= 1:
        return True
    evidence_terms = {
        token.lower()
        for item in list(claim.get("matched_sources") or [])
        if isinstance(item, dict)
        for token in re.findall(
            r"[A-Za-z][A-Za-z0-9_-]{3,}|[\u4e00-\u9fff]{2,}",
            str(item.get("evidence_preview") or ""),
        )
    }
    return all(
        len(
            {
                token.lower()
                for token in re.findall(r"[A-Za-z][A-Za-z0-9_-]{3,}|[\u4e00-\u9fff]{2,}", clause)
            }
            & evidence_terms
        )
        >= 2
        for clause in contrast_clauses
    )


def _matrix_all_claims_fully_bound(verification: dict[str, Any]) -> bool:
    claims = [
        claim
        for claim in list(verification.get("claims") or [])
        if isinstance(claim, dict) and str(claim.get("claim_kind") or "local_claim") == "local_claim"
    ]
    return bool(claims and all(_matrix_claim_fully_bound(claim) for claim in claims))


def _matrix_answer_passes(verification: dict[str, Any], expected_sources: set[str]) -> bool:
    return bool(
        _safe_int(verification.get("total_claims")) > 0
        and _safe_int(verification.get("total_claims")) <= 8
        and _safe_int(verification.get("unsupported_claims")) == 0
        and str(verification.get("evidence_status") or "").lower() == "grounded"
        and _matrix_all_claims_fully_bound(verification)
        and expected_sources.issubset(_matrix_verified_sources(verification))
    )


def _repair_matrix_candidate(
    answer: str,
    hits: list[dict[str, Any]],
    verification: dict[str, Any],
    expected_sources: set[str],
    *,
    prefer_zh: bool,
) -> tuple[str, dict[str, Any]]:
    from kb.agent.tools import verify_answer_citations
    from kb.agent.verifier import classify_answer_claims

    local_rows = [
        row
        for row in list(verification.get("claims") or [])
        if isinstance(row, dict) and str(row.get("claim_kind") or "local_claim") == "local_claim"
    ]
    row_index = 0
    kept_claims: list[str] = []
    seen_claims: set[str] = set()
    for item in classify_answer_claims(answer, answer_mode="evidence_grounded"):
        if item.kind != "local_claim":
            continue
        row = local_rows[row_index] if row_index < len(local_rows) else {}
        row_index += 1
        text = re.sub(r"\s+", " ", str(item.text or "")).strip()
        key = text.lower()
        if _matrix_claim_fully_bound(row) and text and key not in seen_claims:
            kept_claims.append(text)
            seen_claims.add(key)
        if len(kept_claims) >= 8:
            break

    repaired_lines = ["## 证据" if prefer_zh else "## Evidence"]
    repaired_lines.extend(f"- {claim}" for claim in kept_claims)
    interim_answer = "\n".join(repaired_lines).strip()
    interim_verification = verify_answer_citations(
        interim_answer,
        hits,
        answer_mode="evidence_grounded",
    )["verification"]
    missing_sources = expected_sources - _matrix_verified_sources(interim_verification)
    supplemented = 0
    seen_sources: set[str] = set()
    for citation_number, hit in enumerate(hits, start=1):
        source_key = _matrix_source_key(hit)
        if not source_key or source_key in seen_sources or source_key not in missing_sources:
            continue
        seen_sources.add(source_key)
        sentence = _matrix_extractive_sentence(hit, citation_number)
        if sentence:
            repaired_lines.append(f"- {sentence}")
            supplemented += 1
    fully_bound_count = len([row for row in local_rows if _matrix_claim_fully_bound(row)])
    removed_invalid = max(0, len(local_rows) - fully_bound_count)
    removed_out_of_contract = max(0, fully_bound_count - len(kept_claims))
    return "\n".join(repaired_lines).strip(), {
        "attempted": True,
        "candidate_model_claims": len(local_rows),
        "preserved_model_claims": len(kept_claims),
        "removed_claims_total": removed_invalid + removed_out_of_contract,
        "removed_unsupported_claims": removed_invalid,
        "removed_out_of_contract_claims": removed_out_of_contract,
        "removed_strict_gate_claims": len(
            [row for row in local_rows if bool(row.get("supported")) and not _matrix_claim_fully_bound(row)]
        ),
        "supplemented_source_claims": supplemented,
    }


def generate_research_brief_from_matrix(
    prompt: str,
    *,
    matrix_record: dict[str, Any],
    settings: Any,
    max_tokens: int = 1_800,
) -> dict[str, Any]:
    from kb.agent.tools import generate_grounded_answer, verify_answer_citations
    from kb.evidence_matrix import evidence_matrix_hits

    total_started = time.perf_counter()
    plan_started = time.perf_counter()
    hits = evidence_matrix_hits(matrix_record, limit=20)
    rows = [item for item in list(matrix_record.get("rows") or []) if isinstance(item, dict)]
    claim_plan = _matrix_claim_plan(hits)
    plan_ms = round((time.perf_counter() - plan_started) * 1000, 2)
    agent_notes = {
        "answer_contract": "research_brief",
        "evidence_gate": {
            "answer_mode": "evidence_grounded",
            "source_blend": "local_grounded",
            "source_policy": "local_only",
        },
        "evidence_matrix": rows[:8],
        "verified_comparison_audits": [
            item
            for item in list(matrix_record.get("comparison_audits") or [])
            if isinstance(item, dict) and str(item.get("status") or "") == "verified"
        ][:8],
    }
    generation_prompt = _matrix_generation_prompt(prompt, claim_plan)
    generation_started = time.perf_counter()
    generated = generate_grounded_answer(
        generation_prompt,
        hits,
        settings=settings,
        agent_notes=agent_notes,
        temperature=0.1,
        max_tokens=max_tokens,
        defer_quality_gate_repair=True,
    )
    generation_ms = round((time.perf_counter() - generation_started) * 1000, 2)
    answer = str(generated.get("answer") or "").strip()
    gate = generated.get("quality_gate") if isinstance(generated.get("quality_gate"), dict) else {}
    verification_started = time.perf_counter()
    verification_payload = verify_answer_citations(answer, hits, answer_mode="evidence_grounded")
    verification = (
        verification_payload.get("verification")
        if isinstance(verification_payload.get("verification"), dict)
        else {}
    )
    initial_verification_ms = round((time.perf_counter() - verification_started) * 1000, 2)
    expected_sources = {
        _first_text(row, "source_path", "source_name", limit=1_200).replace("\\", "/").lower()
        for row in rows
        if _first_text(row, "source_path", "source_name", limit=1_200)
    }
    repair_info: dict[str, Any] = {
        "attempted": False,
        "candidate_model_claims": 0,
        "preserved_model_claims": 0,
        "removed_claims_total": 0,
        "removed_unsupported_claims": 0,
        "removed_out_of_contract_claims": 0,
        "removed_strict_gate_claims": 0,
        "supplemented_source_claims": 0,
    }
    repair_ms = 0.0
    direct_pass = bool(str(gate.get("status") or "").lower() == "passed" and _matrix_answer_passes(verification, expected_sources))
    if not direct_pass:
        repair_started = time.perf_counter()
        answer, repair_info = _repair_matrix_candidate(
            answer,
            hits,
            verification,
            expected_sources,
            prefer_zh=bool(re.search(r"[\u4e00-\u9fff]", str(prompt or ""))),
        )
        verification_payload = verify_answer_citations(answer, hits, answer_mode="evidence_grounded")
        verification = (
            verification_payload.get("verification")
            if isinstance(verification_payload.get("verification"), dict)
            else {}
        )
        repair_ms = round((time.perf_counter() - repair_started) * 1000, 2)
        if (
            generated.get("llm_used") is True
            and _safe_int(repair_info.get("preserved_model_claims")) > 0
            and _matrix_answer_passes(verification, expected_sources)
        ):
            gate = {
                "status": "repaired",
                "reasons": list(dict.fromkeys(list(gate.get("reasons") or []) + ["matrix_claim_repair"])),
                "warnings": ["targeted_claim_repair"],
            }
        else:
            gate = {"status": "fallback", "reasons": ["matrix_source_coverage_or_support_gate"], "warnings": []}
    if str(gate.get("status") or "").lower() == "fallback":
        answer = _matrix_extractive_brief_answer(prompt, hits)
        gate = {
            "status": "fallback",
            "reasons": ["matrix_source_coverage_or_support_gate"],
            "warnings": ["extractive_fallback"],
        }
        verification_payload = verify_answer_citations(answer, hits, answer_mode="evidence_grounded")
        verification = (
            verification_payload.get("verification")
            if isinstance(verification_payload.get("verification"), dict)
            else {}
        )
    errors: list[str] = []
    total_ms = round((time.perf_counter() - total_started) * 1000, 2)
    summary = {
        "query_scope": "basket",
        "quality_gate_status": str(gate.get("status") or ""),
        "quality_gate_reasons": list(gate.get("reasons") or []),
        "quality_gate_warnings": list(gate.get("warnings") or []),
        "matrix_id": str(matrix_record.get("id") or ""),
        "matrix_revision": int(matrix_record.get("revision") or 1),
        "verified_comparison_count": len(
            [
                item
                for item in list(matrix_record.get("comparison_audits") or [])
                if isinstance(item, dict) and str(item.get("status") or "") == "verified"
            ]
        ),
        "claim_repair": repair_info,
        "phase_timings_ms": {
            "claim_plan": plan_ms,
            "model_synthesis": generation_ms,
            "initial_verification": initial_verification_ms,
            "targeted_claim_repair": repair_ms,
            "total": total_ms,
        },
        **{
            key: verification.get(key)
            for key in (
                "total_claims",
                "supported_claims",
                "unsupported_claims",
                "support_ratio",
                "evidence_status",
            )
        },
    }
    trace = {
        "mode": "research_agent",
        "question_type": "multi_paper_comparison",
        "context": {
            "query_scope": "basket",
            "answer_contract": "research_brief",
            "source_matrix_id": str(matrix_record.get("id") or ""),
            "source_matrix_revision": int(matrix_record.get("revision") or 1),
        },
        "plan": claim_plan,
        "steps": [
            {
                "tool": "generate_grounded_answer",
                "status": "done",
                "observation": "Generated the brief only from the verified project evidence matrix.",
                "output": {"quality_gate": gate},
                "error": "",
                "elapsed_ms": int(round(generation_ms)),
            },
            *(
                [
                    {
                        "tool": "repair_matrix_claims",
                        "status": "done",
                        "observation": "Preserved supported model claims and supplemented only missing matrix sources.",
                        "output": repair_info,
                        "error": "",
                        "elapsed_ms": int(round(repair_ms)),
                    }
                ]
                if bool(repair_info.get("attempted"))
                else []
            ),
            {
                "tool": "verify_answer_citations",
                "status": "done",
                "observation": str(verification_payload.get("observation") or ""),
                "output": verification,
                "error": "",
                "elapsed_ms": int(round(initial_verification_ms)),
            },
        ],
        "verification": verification,
        "research_run": {
            "run_id": f"matrix_{str(matrix_record.get('id') or '')[:16]}",
            "status": "verified" if not errors else "failed",
            "source_policy": "local_only",
            "query_scope": "basket",
            "question": _text(prompt, limit=500),
            "subtasks": [],
            "evidence_matrix": rows[:8],
            "metrics": {
                "evidence_matrix_rows": len(rows),
                "source_count": len(expected_sources),
                "local_evidence_hit_count": len(hits),
                "model_synthesis_ms": generation_ms,
                "targeted_claim_repair_ms": repair_ms,
                "total_ms": total_ms,
            },
        },
        "status": "done" if not errors else "failed",
        "errors": errors,
        "summary": summary,
    }
    return {"answer": answer, "hits": hits, "agent_trace": trace}


def _reference_line(item: dict[str, Any], index: int) -> str:
    authors = _text(item.get("authors"), limit=800)
    year = _text(item.get("year"), limit=40)
    title = _text(item.get("title"), limit=800) or f"Reference {index}"
    venue = _text(item.get("venue"), limit=500)
    doi = _text(item.get("doi"), limit=400)
    prefix = f"{authors}. " if authors else ""
    suffix = ". ".join(part for part in (year, venue) if part)
    doi_part = f" DOI: {doi}." if doi else ""
    return f"{prefix}{title}. {suffix}{'.' if suffix else ''}{doi_part}".strip()


def research_brief_markdown(record: dict[str, Any]) -> str:
    title = _text(record.get("title"), limit=240) or "Research brief"
    objective = _text(record.get("objective"), limit=4_000)
    content = _text(record.get("content_markdown"), limit=160_000)
    quality_status = _text(record.get("quality_status"), limit=40) or "draft"
    revision = max(1, int(record.get("revision") or 1))
    lines = [f"# {title}", "", f"> Evidence status: {quality_status}; revision: {revision}."]
    if objective:
        lines.extend(["", "## Research objective", "", objective])
    if content:
        lines.extend(["", content])
    evidence = [item for item in list(record.get("evidence") or []) if isinstance(item, dict)]
    if evidence:
        lines.extend(["", "## Evidence appendix"])
        for item in evidence:
            number = int(item.get("citation_number") or 0)
            source = _text(item.get("source_name") or item.get("source_path"), limit=500) or "Source"
            locator = _text(item.get("heading_path") or item.get("location_label"), limit=800)
            quote = _text(item.get("evidence_quote"), limit=_MAX_EVIDENCE_TEXT)
            label = f"[{number}] {source}" if number > 0 else source
            lines.extend(["", f"### {label}{f' — {locator}' if locator else ''}"])
            if quote:
                lines.extend(["", f"> {quote.replace(chr(10), chr(10) + '> ')}"])
    bibliography = [item for item in list(record.get("bibliography") or []) if isinstance(item, dict)]
    if bibliography:
        lines.extend(["", "## References", ""])
        lines.extend(f"{index}. {_reference_line(item, index)}" for index, item in enumerate(bibliography, start=1))
    return "\n".join(lines).strip() + "\n"


def _bibtex_key(item: dict[str, Any], index: int) -> str:
    authors = _text(item.get("authors"), limit=200)
    author = re.sub(r"[^A-Za-z0-9]+", "", authors.split(",")[0].split()[-1] if authors else "")
    year = re.sub(r"\D+", "", _text(item.get("year"), limit=20))
    title = re.sub(r"[^A-Za-z0-9]+", "", _text(item.get("title"), limit=120).split(" ")[0])
    return f"{author or 'source'}{year or 'nd'}{title or index}"[:80]


def _bibtex_value(value: object) -> str:
    return _text(value, limit=2_000).replace("{", "\\{").replace("}", "\\}")


def research_brief_bibtex(record: dict[str, Any]) -> str:
    blocks: list[str] = []
    for index, item in enumerate(list(record.get("bibliography") or []), start=1):
        if not isinstance(item, dict):
            continue
        fields = {
            "title": item.get("title"),
            "author": item.get("authors"),
            "year": item.get("year"),
            "journal": item.get("venue"),
            "doi": item.get("doi"),
        }
        body = ",\n".join(
            f"  {key} = {{{_bibtex_value(value)}}}"
            for key, value in fields.items()
            if _text(value, limit=2_000)
        )
        blocks.append(f"@article{{{_bibtex_key(item, index)},\n{body}\n}}")
    return "\n\n".join(blocks).strip() + ("\n" if blocks else "")


def research_brief_ris(record: dict[str, Any]) -> str:
    blocks: list[str] = []
    for item in list(record.get("bibliography") or []):
        if not isinstance(item, dict):
            continue
        lines = ["TY  - JOUR"]
        authors = _text(item.get("authors"), limit=800)
        for author in [part.strip() for part in re.split(r";|\band\b", authors) if part.strip()]:
            lines.append(f"AU  - {author}")
        for tag, key in (("TI", "title"), ("PY", "year"), ("JO", "venue"), ("DO", "doi")):
            value = _text(item.get(key), limit=1_000)
            if value:
                lines.append(f"{tag}  - {value}")
        lines.append("ER  -")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks).strip() + ("\n" if blocks else "")


def research_brief_docx(record: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(_text(record.get("title"), limit=240) or "Research brief", level=0)
    objective = _text(record.get("objective"), limit=4_000)
    if objective:
        document.add_heading("Research objective", level=1)
        document.add_paragraph(objective)
    content = _text(record.get("content_markdown"), limit=160_000)
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            document.add_heading(heading.group(2).strip(), level=min(4, len(heading.group(1))))
            continue
        if re.match(r"^[-*+]\s+", line):
            document.add_paragraph(re.sub(r"^[-*+]\s+", "", line), style="List Bullet")
            continue
        if re.match(r"^\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", line), style="List Number")
            continue
        document.add_paragraph(line.lstrip("> "))
    evidence = [item for item in list(record.get("evidence") or []) if isinstance(item, dict)]
    if evidence:
        document.add_heading("Evidence appendix", level=1)
        for item in evidence:
            number = int(item.get("citation_number") or 0)
            source = _text(item.get("source_name") or item.get("source_path"), limit=500) or "Source"
            locator = _text(item.get("heading_path") or item.get("location_label"), limit=800)
            document.add_heading(f"[{number}] {source}" if number > 0 else source, level=2)
            if locator:
                document.add_paragraph(locator)
            quote = _text(item.get("evidence_quote"), limit=_MAX_EVIDENCE_TEXT)
            if quote:
                document.add_paragraph(quote)
    bibliography = [item for item in list(record.get("bibliography") or []) if isinstance(item, dict)]
    if bibliography:
        document.add_heading("References", level=1)
        for index, item in enumerate(bibliography, start=1):
            document.add_paragraph(_reference_line(item, index), style="List Number")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
