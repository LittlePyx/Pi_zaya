from __future__ import annotations

from io import BytesIO

from docx import Document

from kb.research_note import research_note_docx, research_note_filename


def test_research_note_filename_preserves_unicode_and_removes_windows_separators() -> None:
    assert research_note_filename("  单光子：结果 / 对比？  ") == "单光子：结果 - 对比？.docx"
    assert research_note_filename("<>|\x00") == "Pi_zaya research note.docx"


def test_research_note_docx_preserves_core_markdown_structure():
    payload = research_note_docx(
        "单光子成像研究笔记",
        """## 研究问题 1

为什么需要建模噪声？

### 回答

结论包含 **加粗内容**、`参数` 与 [DOI 链接](https://doi.org/10.1000/example)。

> 可定位的原文依据。

| 方法 | PSNR |
| --- | ---: |
| Ours | 22.36 |

1. First source.
""",
    )

    document = Document(BytesIO(payload))
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert document.core_properties.title == "单光子成像研究笔记"
    assert "研究问题 1" in all_text
    assert "可定位的原文依据" in all_text
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "方法"
    assert document.tables[0].cell(1, 1).text == "22.36"
    assert any(
        relationship.target_ref == "https://doi.org/10.1000/example"
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink")
    )


def test_research_note_docx_keeps_unclosed_code_fence_content():
    payload = research_note_docx("Code", "```python\nvalue = 3\nprint(value)")
    document = Document(BytesIO(payload))
    assert any("value = 3" in paragraph.text for paragraph in document.paragraphs)


def test_research_note_docx_removes_null_bytes_and_uses_fallback_title():
    payload = research_note_docx("\x00", "正文\x00内容")
    document = Document(BytesIO(payload))
    assert document.core_properties.title == "Pi_zaya research note"
    assert any("正文内容" in paragraph.text for paragraph in document.paragraphs)
